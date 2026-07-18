"""
scheduler.py - everything around the algorithms: config/state I/O, the
calendar, human-readable reporting, independent validation, and file export
(CSV/DOCX/PDF). Nothing in here decides WHO does WHAT - that's optimizer.py.
This module just loads inputs, reports outputs, and exports them.
"""
import sys
import datetime
import json
import os
import re
import yaml

from optimizer import (
    is_person_off, build_windows, resolve_group_windows, buffer_threshold,
    evaluate, polish_schedule, count_structural_violations, algo_milp,
    rest_balance_spread,
)


def load_config(path="config.yml"):
    try:
        with open(path, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f)
    except FileNotFoundError:
        print(f"🚨 Config file not found: {path}")
        sys.exit()
    for person in cfg["roommates"]:
        cfg["days_off"].setdefault(person, [])
        if cfg["days_off"][person] is None:
            cfg["days_off"][person] = []

    by_name = {g["name"]: g for g in cfg["chore_groups"]}
    for g in cfg["chore_groups"]:
        if g.get("buffer_days") is None:
            g["buffer_days"] = cfg["buffer_days"]
        if g.get("tolerance_days") is None:
            g["tolerance_days"] = 0
        if "piggyback_on" in g:
            if g["piggyback_on"] not in by_name:
                print(f"🚨 chore group '{g['name']}' piggybacks on unknown group '{g['piggyback_on']}'")
                sys.exit()
            if not g.get("every_nth") or g["every_nth"] < 1:
                print(f"🚨 chore group '{g['name']}' needs every_nth >= 1 alongside piggyback_on")
                sys.exit()
            # approximate frequency for greedy/SA/feasibility-check, which don't
            # understand piggyback alignment - only the MILP enforces it exactly
            host = by_name[g["piggyback_on"]]
            g["frequency_days"] = host["frequency_days"] * g["every_nth"]
        elif not g.get("frequency_days"):
            print(f"🚨 chore group '{g['name']}' needs either frequency_days or piggyback_on+every_nth")
            sys.exit()
    return cfg




def build_calendar(start_date, total_days):
    calendar = [start_date + datetime.timedelta(days=d) for d in range(total_days)]
    weekday = [dt.strftime("%A") for dt in calendar]
    return calendar, weekday




def load_state(path="schedule_state.json"):
    """State from the previous run: last task date+chore per person (for rest
    carry-over), last occurrence date per chore group (for phase continuity),
    and cumulative task counts (for fairness across multiple runs)."""
    import json
    import os
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    raw["people"] = {
        p: {"last_task_date": datetime.datetime.strptime(v["last_task_date"], "%Y-%m-%d").date(),
            "last_group": v["last_group"]}
        for p, v in raw.get("people", {}).items()
    }
    raw["group_phase"] = {
        g: datetime.datetime.strptime(v, "%Y-%m-%d").date()
        for g, v in raw.get("group_phase", {}).items()
    }
    return raw




def save_state(winner_slots, people, chore_groups, path="schedule_state.json"):
    """Persist exactly what the NEXT run needs to continue seamlessly, no more."""
    import json
    state = {"people": {}, "group_phase": {}, "cumulative_totals": {}, "cumulative_per_task": {}}

    by_person = {p: [] for p in people}
    for s in winner_slots:
        by_person[s["person"]].append(s)
    for p in people:
        slots = sorted(by_person[p], key=lambda s: s["day_idx"])
        if slots:
            last = slots[-1]
            state["people"][p] = {
                "last_task_date": last["date"].strftime("%Y-%m-%d"),
                "last_group": chore_groups[last["group_idx"]]["name"],
            }
        state["cumulative_totals"][p] = len(slots)

    for g in chore_groups:
        g_slots = [s for s in winner_slots if chore_groups[s["group_idx"]]["name"] == g["name"]]
        if g_slots:
            last_date = max(s["date"] for s in g_slots)
            state["group_phase"][g["name"]] = last_date.strftime("%Y-%m-%d")

    from collections import defaultdict
    per_task = defaultdict(lambda: defaultdict(int))
    for s in winner_slots:
        per_task[s["task"]][s["person"]] += 1
    state["cumulative_per_task"] = {t: dict(counts) for t, counts in per_task.items()}

    with open(path, "w", encoding="utf-8") as f:
        json.dump(state, f, indent=2)
    return path


def diagnose_milp_infeasibility(roommates, chore_groups, days_off, calendar, weekday, state):
    """Runs automatically when MILP goes infeasible on a continuation run.
    Isolates whether rest carry-over, phase carry-over, or fairness
    carry-over (or some combination) is actually responsible, by re-solving
    with each turned off one at a time. This is the exact manual process
    used to debug this class of issue before - now it happens automatically,
    in your own terminal, without needing to hand over schedule_state.json."""
    import copy

    print("=" * 70)
    print(" DIAGNOSING MILP INFEASIBILITY (continuation run)")
    print(" Testing which carry-over component is responsible...")
    print("=" * 70)

    tests = [("no carry-over at all (state=None)", None)]
    if state.get("people"):
        s = copy.deepcopy(state)
        s["group_phase"], s["cumulative_totals"], s["cumulative_per_task"] = {}, {}, {}
        tests.append(("rest carry-over ONLY", s))
    if state.get("group_phase"):
        s = copy.deepcopy(state)
        s["people"], s["cumulative_totals"], s["cumulative_per_task"] = {}, {}, {}
        tests.append(("phase (chore cadence) carry-over ONLY", s))
    if state.get("cumulative_totals") or state.get("cumulative_per_task"):
        s = copy.deepcopy(state)
        s["people"], s["group_phase"] = {}, {}
        tests.append(("fairness carry-over ONLY", s))
    tests.append(("ALL THREE carry-overs (what actually ran)", state))

    results = []
    for name, test_state in tests:
        slots, note, fv, cv = algo_milp(roommates, chore_groups, days_off, calendar, weekday, state=test_state)
        status = "✅ FEASIBLE" if slots else "❌ INFEASIBLE"
        results.append((name, status))
        print(f"  {status}  <-  {name}")

    print("=" * 70)
    culprits = [name for name, status in results if "INFEASIBLE" in status and name != "ALL THREE carry-overs (what actually ran)"]
    if culprits:
        print(f" LIKELY CAUSE: {', '.join(culprits)}")
        print(" This carry-over, combined with your current days-off/frequencies,")
        print(" creates a genuinely unsatisfiable hard constraint. Consider:")
        print("   - loosening that specific carry-over (e.g. skip it just this once)")
        print("   - adjusting the config that's in conflict (buffer_days, frequency, days_off)")
    else:
        print(" No single component alone is infeasible - it's the COMBINATION.")
        print(" This is rarer and harder to loosen selectively; the safest fix is")
        print(" accepting the best-effort heuristic winner for this one run.")
    print("=" * 70)
    return results


def polish_and_report(name, slots, roommates, chore_groups, days_off, calendar, weekday, buffer_days_target):
    """Apply the rest-polish layer to one algorithm's output and print a
    compact before -> after summary. Also re-checks structural correctness
    before/after as a safety net - polish_schedule already guards against
    increasing structural violations internally, but it's cheap to verify
    rather than assume."""
    if slots is None:
        return None
    before = evaluate(roommates, chore_groups, slots)
    before_structural = count_structural_violations(slots, chore_groups, days_off, calendar, weekday)

    polished, moves = polish_schedule(roommates, chore_groups, days_off, slots, buffer_days_target, calendar, weekday)

    after = evaluate(roommates, chore_groups, polished)
    after_structural = count_structural_violations(polished, chore_groups, days_off, calendar, weekday)

    print(f"   Polish: {moves} move(s) -> "
          f"avg rest {before['avg_rest']:.2f}->{after['avg_rest']:.2f}, "
          f"min rest {before['min_rest']}->{after['min_rest']}, "
          f"violations {before['buffer_violations']}->{after['buffer_violations']}")
    if after_structural > before_structural:
        print(f"   🚨 unexpected: structural violations INCREASED during polish "
              f"({before_structural}->{after_structural}) - this should never happen, keeping original")
        return slots
    return polished




def print_metrics(name, active_slots, people, chore_groups, elapsed, note):
    if active_slots is None:
        print(f"\n### {name}: 🚨 INFEASIBLE ({note})")
        return
    m = evaluate(people, chore_groups, active_slots)
    print(f"\n### {name}  [{note}]  ({elapsed*1000:.1f} ms)")
    print(f"   Fairness spread (max-min tasks) : {m['fairness_spread']}")
    print(f"   Workload : " + ", ".join(f"{p}={c}" for p, c in m['workload'].items()))
    print(f"   Avg / Min / Max rest gap        : {m['avg_rest']:.2f} / {m['min_rest']} / {m['max_rest']}")
    print(f"   Buffer rule violations          : {m['buffer_violations']}")
    print(f"   Per-chore fairness (sum spread) : {m['total_per_task_spread']}  "
          + ", ".join(f"{t}:{s}" for t, s in m["per_task_spread"].items()))

    last_task = {}
    per_person_rests = {p: [] for p in people}
    for slot in sorted(active_slots, key=lambda s: s["day_idx"]):
        p = slot["person"]
        if p in last_task:
            per_person_rests[p].append(slot["day_idx"] - last_task[p] - 1)
        last_task[p] = slot["day_idx"]
    per_person_avgs = {p: (sum(rs) / len(rs) if rs else None) for p, rs in per_person_rests.items()}
    spread = rest_balance_spread(people, active_slots)
    print(f"   Rest balance across people      : spread {spread:.2f}  ("
          + ", ".join(f"{p}={v:.2f}" if v is not None else f"{p}=n/a" for p, v in per_person_avgs.items()) + ")")




def print_task_breakdown(active_slots, chore_groups):
    """Per-chore-type breakdown: who does THIS specific chore, and how often."""
    from collections import defaultdict
    by_task = defaultdict(lambda: defaultdict(int))
    for s in active_slots:
        by_task[s["task"]][s["person"]] += 1

    # keep chore_groups' declared order/task order for readability
    ordered_tasks = [t for g in chore_groups for t in g["tasks"]]

    print("-" * 70)
    print(" TASK BREAKDOWN BY PERSON")
    print("-" * 70)
    for task in ordered_tasks:
        counts = by_task.get(task, {})
        if not counts:
            continue
        parts = ", ".join(f"{p}({c})" for p, c in sorted(counts.items(), key=lambda kv: -kv[1]))
        spread = max(counts.values()) - min(counts.values()) if len(counts) > 1 else 0
        flag = "  ⚠️ uneven" if spread >= 2 else ""
        print(f"  {task:18s} : {parts}{flag}")
    print()




def print_person_summary(active_slots, people):
    """Per-person: total task count, and the full sequence of rest gaps."""
    by_person = {p: [] for p in people}
    for s in active_slots:
        by_person[s["person"]].append(s)
    for p in by_person:
        by_person[p].sort(key=lambda s: s["day_idx"])

    print("-" * 70)
    print(" TOTAL TASKS BY PERSON")
    print("-" * 70)
    for p in people:
        print(f"  {p:10s} : {len(by_person[p])} tasks")
    print()

    print("-" * 70)
    print(" REST PERIODS BY PERSON (days between consecutive tasks)")
    print("-" * 70)
    for p in people:
        slots = by_person[p]
        rests = []
        for i in range(1, len(slots)):
            gap = slots[i]["day_idx"] - slots[i - 1]["day_idx"]
            rests.append(gap - 1)
        if rests:
            avg_r = sum(rests) / len(rests)
            print(f"  {p:10s} : {rests}   (avg {avg_r:.1f}, min {min(rests)}, max {max(rests)})")
        else:
            print(f"  {p:10s} : no rest data (0-1 tasks total)")
    print()




def validate_schedule(active_slots, chore_groups, days_off, calendar, weekday):
    """Independent audit: re-derive every rule from config and check the
    ACTUAL schedule against it. This does not trust the solver/heuristic -
    it re-checks from scratch, the same way a human would verify by hand."""
    total_days = len(calendar)
    chore_groups_by_name = {g["name"]: g for g in chore_groups}
    occ_by_group = {g["name"]: sorted({s["day_idx"] for s in active_slots
                                        if chore_groups[s["group_idx"]]["name"] == g["name"]})
                    for g in chore_groups}

    print("=" * 70)
    print(" SCHEDULE VALIDATION")
    print("=" * 70)
    all_ok = True

    # --- 1) frequency check: every non-piggyback group fires EXACTLY once
    # per its own frequency window (this covers Dishes freq=1, Stove freq=7,
    # SweepTrash freq=3 all with the same generic check) ---
    for g in chore_groups:
        if "piggyback_on" in g:
            continue
        windows = build_windows(g["frequency_days"], total_days)
        occ = set(occ_by_group[g["name"]])
        bad_windows = [w for w in windows if len(occ & set(w)) != 1]
        status = "OK" if not bad_windows else "FAIL"
        if bad_windows:
            all_ok = False
        print(f"  [{status}] {g['name']} (every {g['frequency_days']}d): "
              f"{len(windows)} windows, {len(occ)} occurrences"
              + (f" - bad windows: {[[calendar[d].isoformat() for d in w] for w in bad_windows]}" if bad_windows else ""))

        # --- 1b) cadence check: consecutive occurrences should stay close to
        # frequency_days apart, not just "one per independent window" -
        # catches e.g. two Stove visits 12 days apart when freq=7.
        # tolerance_days relaxes the minimum acceptable gap. ---
        tol = g.get("tolerance_days", 0)
        effective_min_gap = max(1, g["frequency_days"] - tol)
        occ_sorted = sorted(occ)
        bad_gaps = []
        for i in range(1, len(occ_sorted)):
            gap = occ_sorted[i] - occ_sorted[i - 1]
            if gap < effective_min_gap or gap > 2 * g["frequency_days"] - 1 + tol:
                bad_gaps.append((calendar[occ_sorted[i - 1]].isoformat(), calendar[occ_sorted[i]].isoformat(), gap))
        cadence_status = "OK" if not bad_gaps else "FAIL"
        if bad_gaps:
            all_ok = False
        tol_note = f" (tolerance +-{tol}d)" if tol else ""
        print(f"  [{cadence_status}] {g['name']} cadence (expect ~{g['frequency_days']}d apart{tol_note}): "
              + (f"irregular gaps: {bad_gaps}" if bad_gaps else "consistent"))

    # --- 2) piggyback check: every piggyback occurrence must land on an
    # occurrence day of its host (exactly, or within tolerance_days) ---
    for g in chore_groups:
        if "piggyback_on" not in g:
            continue
        host_occ = occ_by_group[g["piggyback_on"]]
        every_nth = g["every_nth"]
        tol = g.get("tolerance_days", 0)
        expected = [d for i, d in enumerate(host_occ) if (i + 1) % every_nth == 0]
        actual = occ_by_group[g["name"]]
        if len(actual) != len(expected):
            status, mismatches = "FAIL", None
        else:
            mismatches = [(a, e) for a, e in zip(actual, expected) if abs(a - e) > tol]
            status = "OK" if not mismatches else "FAIL"
        if status == "FAIL":
            all_ok = False
        exp_dates = [calendar[d].isoformat() for d in expected]
        act_dates = [calendar[d].isoformat() for d in actual]
        tol_note = f" (tolerance +-{tol}d)" if tol else ""
        print(f"  [{status}] {g['name']} piggybacks on {g['piggyback_on']} every {every_nth}{tol_note}: "
              f"expected {exp_dates}, got {act_dates}")

    # --- 3) days-off conflict check ---
    conflicts = []
    for s in active_slots:
        if is_person_off(s["person"], s["date"], s["day"], days_off):
            conflicts.append(s)
    status = "OK" if not conflicts else "FAIL"
    if conflicts:
        all_ok = False
    print(f"  [{status}] Days-off conflicts: {len(conflicts)}"
          + (f" - {[(c['person'], c['date'].isoformat(), c['task']) for c in conflicts]}" if conflicts else ""))

    print("=" * 70)
    print(f" VALIDATION {'PASSED' if all_ok else 'FAILED'}")
    print("=" * 70)
    return all_ok




def print_schedule(name, active_slots, chore_groups):
    """Full week-by-week timeline for one algorithm's result."""
    last_task = {}
    rows = []
    for slot in sorted(active_slots, key=lambda s: s["day_idx"]):
        p = slot["person"]
        if p in last_task:
            rest = slot["day_idx"] - last_task[p] - 1
        else:
            rest = None
        rows.append({**slot, "rest": rest})
        last_task[p] = slot["day_idx"]

    days_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    weekly = {}
    for r in rows:
        monday = r["date"] - datetime.timedelta(days=r["date"].weekday())
        week_str = f"Week of {monday.strftime('%b %d, %Y')}"
        weekly.setdefault(week_str, {d: [] for d in days_order})
        rest_str = f"[Rest: {r['rest']}]" if r["rest"] is not None else "[First]"
        weekly[week_str][r["day"]].append(f"{r['task']}: {r['person']} {rest_str}")

    print("\n" + "#" * 70)
    print(f"# FULL SCHEDULE - {name}")
    print("#" * 70)
    for week, days in weekly.items():
        print("=" * 100)
        print(f" 🗓️  {week.upper()} ")
        print("=" * 100)
        for day in days_order:
            if days[day]:
                print(f"{day:<12} | {'  |  '.join(days[day])}")
    print()




def schedule_rows_for_export(active_slots):
    """Flat, sorted list of rows shared by all three export formats."""
    last_task = {}
    rows = []
    for slot in sorted(active_slots, key=lambda s: s["day_idx"]):
        p = slot["person"]
        rest = (slot["day_idx"] - last_task[p] - 1) if p in last_task else None
        rows.append({
            "date": slot["date"].strftime("%Y-%m-%d"),
            "day": slot["day"],
            "task": slot["task"],
            "person": p,
            "rest": "First" if rest is None else str(rest),
        })
        last_task[p] = slot["day_idx"]
    return rows




def export_csv(active_slots, path="chore_schedule.csv"):
    import csv
    rows = schedule_rows_for_export(active_slots)
    with open(path, "w", newline="", encoding="utf-8-sig") as f:  # utf-8-sig so Excel shows emoji/accents correctly
        writer = csv.DictWriter(f, fieldnames=["date", "day", "task", "person", "rest"])
        writer.writerow({"date": "Date", "day": "Day", "task": "Task", "person": "Assigned To", "rest": "Rest Taken"})
        writer.writerows(rows)
    return path




def export_docx(active_slots, path="chore_schedule.docx"):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rows = schedule_rows_for_export(active_slots)
    doc = Document()
    doc.add_heading("Chore Schedule", level=1)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Date", "Day", "Task", "Assigned To", "Rest Taken"]):
        hdr[i].text = label
        hdr[i].paragraphs[0].runs[0].bold = True

    for r in rows:
        cells = table.add_row().cells
        cells[0].text = r["date"]
        cells[1].text = r["day"]
        cells[2].text = r["task"]
        cells[3].text = r["person"]
        cells[4].text = r["rest"]

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.save(path)
    return path




def export_pdf(active_slots, chore_groups, people, path="chore_schedule.pdf"):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle, Paragraph,
                                     Spacer, PageBreak)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    import re

    def strip_emoji(text):
        # ReportLab's built-in fonts can't render emoji glyphs (blank boxes) -
        # the descriptive word after it (e.g. "Dishes") already carries the meaning.
        return re.sub(r"[^\x00-\x7F]+", "", text).strip()

    # a fixed color per chore type, used consistently across every page -
    # scan the calendar by color instead of reading every word
    palette = ["#4472C4", "#C0504D", "#9BBB59", "#8064A2", "#4BACC6", "#F79646"]
    ordered_tasks = [t for g in chore_groups for t in g["tasks"]]
    task_color = {strip_emoji(t): palette[i % len(palette)] for i, t in enumerate(ordered_tasks)}

    styles = getSampleStyleSheet()
    cell_style = ParagraphStyle("cell", parent=styles["Normal"], fontSize=8, leading=10)
    day_header_style = ParagraphStyle("dayhdr", parent=styles["Normal"], fontSize=9,
                                       textColor=colors.white, alignment=TA_CENTER, fontName="Helvetica-Bold")

    rows = schedule_rows_for_export(active_slots)
    days_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

    # group rows by week, then by day-of-week
    weekly = {}
    week_mondays = {}
    for r in rows:
        d = datetime.datetime.strptime(r["date"], "%Y-%m-%d").date()
        monday = d - datetime.timedelta(days=d.weekday())
        week_key = monday.strftime("%b %d, %Y")
        weekly.setdefault(week_key, {day: [] for day in days_order})
        week_mondays[week_key] = monday
        weekly[week_key][r["day"]].append(r)

    doc = SimpleDocTemplate(path, pagesize=letter,
                             leftMargin=0.4 * inch, rightMargin=0.4 * inch,
                             topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    elements = [Paragraph("Chore Schedule", styles["Title"]), Spacer(1, 10)]

    # --- intro block: schedule period, roommates, chores + frequency, totals ---
    all_dates = sorted(datetime.datetime.strptime(r["date"], "%Y-%m-%d").date() for r in rows)
    start_date_str = all_dates[0].strftime("%b %d, %Y") if all_dates else "-"
    end_date_str = all_dates[-1].strftime("%b %d, %Y") if all_dates else "-"
    n_weeks = ((all_dates[-1] - all_dates[0]).days // 7 + 1) if all_dates else 0

    def chore_freq_label(g):
        tol = g.get("tolerance_days", 0)
        tol_note = f" (+-{tol}d)" if tol else ""
        if "piggyback_on" in g:
            return f"{strip_emoji(g['tasks'][0])} - piggybacks on {g['piggyback_on']} every {g['every_nth']}{tol_note}"
        return f"{strip_emoji(g['tasks'][0])} - every {g['frequency_days']} day(s){tol_note}"

    buffer_values = {g.get("buffer_days", 0) for g in chore_groups}
    buffer_label = f"{buffer_values.pop()} days" if len(buffer_values) == 1 else "varies by chore (see config)"

    intro_data = [
        ["Schedule period", f"{start_date_str}  -  {end_date_str}  ({n_weeks} week{'s' if n_weeks != 1 else ''})"],
        [f"Roommates ({len(people)})", ", ".join(people)],
        ["Chores and frequency", "\n".join(chore_freq_label(g) for g in chore_groups)],
        ["Total tasks scheduled", str(len(rows))],
        ["Rest buffer target", buffer_label],
    ]
    intro_rows_formatted = []
    for label, value in intro_data:
        intro_rows_formatted.append([
            Paragraph(f"<b>{label}</b>", cell_style),
            Paragraph(value.replace(chr(10), "<br/>"), cell_style),
        ])
    intro_table = Table(intro_rows_formatted, colWidths=[1.6 * inch, 5.2 * inch])
    intro_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F2F2")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    elements += [intro_table, Spacer(1, 16)]

    # --- summary page: total tasks per person + color legend ---
    task_counts = {p: 0 for p in people}
    for r in rows:
        task_counts[r["person"]] += 1
    summary_data = [["Person", "Total Tasks"]] + [[p, str(c)] for p, c in task_counts.items()]
    summary_table = Table(summary_data, colWidths=[2 * inch, 1.2 * inch])
    summary_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
    ]))
    elements += [Paragraph("Total tasks by person", styles["Heading3"]), summary_table, Spacer(1, 16)]

    # --- task breakdown by person (per chore type) - same data as the console
    # report, but clean: no "uneven" warning annotations in the PDF ---
    from collections import defaultdict
    by_task = defaultdict(lambda: defaultdict(int))
    for r in rows:
        by_task[strip_emoji(r["task"])][r["person"]] += 1
    breakdown_data = [["Chore", "Breakdown"]]
    for task in ordered_tasks:
        task_clean = strip_emoji(task)
        counts = by_task.get(task_clean, {})
        if not counts:
            continue
        parts = ", ".join(f"{p}({c})" for p, c in sorted(counts.items(), key=lambda kv: -kv[1]))
        breakdown_data.append([task_clean, parts])
    breakdown_table = Table(breakdown_data, colWidths=[1.3 * inch, 5.5 * inch])
    breakdown_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements += [Paragraph("Task breakdown by person", styles["Heading3"]), breakdown_table, Spacer(1, 16)]

    # --- rest periods by person - same data as the console report, clean (no flags) ---
    by_person_slots = {p: [] for p in people}
    for r in rows:
        by_person_slots[r["person"]].append(r)
    rest_data = [["Person", "Rest gaps (days between tasks)", "Avg", "Min", "Max"]]
    for p in people:
        p_rows = sorted(by_person_slots[p], key=lambda r: r["date"])
        rests = []
        for i in range(1, len(p_rows)):
            d1 = datetime.datetime.strptime(p_rows[i - 1]["date"], "%Y-%m-%d").date()
            d2 = datetime.datetime.strptime(p_rows[i]["date"], "%Y-%m-%d").date()
            rests.append((d2 - d1).days - 1)
        if rests:
            rest_data.append([p, ", ".join(str(r) for r in rests),
                              f"{sum(rests) / len(rests):.1f}", str(min(rests)), str(max(rests))])
        else:
            rest_data.append([p, "no rest data (0-1 tasks total)", "-", "-", "-"])
    rest_table = Table(rest_data, colWidths=[0.9 * inch, 4.4 * inch, 0.5 * inch, 0.5 * inch, 0.5 * inch])
    rest_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    elements += [Paragraph("Rest periods by person", styles["Heading3"]), rest_table, Spacer(1, 16)]

    legend_cells = []
    for task, hexcolor in task_color.items():
        legend_cells.append(Paragraph(f'<font color="{hexcolor}">\u25A0</font> {task}', cell_style))
    legend_table = Table([legend_cells])
    legend_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    elements += [Paragraph("Chore color key", styles["Heading3"]), legend_table, PageBreak()]

    # --- one PERSON TIMELINE page per week: rows = people, columns = days ---
    # answers "when am I on duty" at a glance, and shows task-pairing patterns
    # (e.g. Mop riding on a Sweep day) as a visible row pattern, not buried text
    name_col_width = 0.9 * inch
    day_col_width = (7.3 * inch - name_col_width) / 7

    for week_idx, (week_key, days) in enumerate(weekly.items()):
        if week_idx > 0:
            elements.append(PageBreak())
        elements.append(Paragraph(f"Week of {week_key}", styles["Heading2"]))
        elements.append(Spacer(1, 6))

        monday_date = week_mondays[week_key]
        header_row = [Paragraph("", day_header_style)] + [
            Paragraph(f"{day[:3]}<br/><font size=\"7\">{(monday_date + datetime.timedelta(days=idx)).strftime('%b %d')}</font>",
                      day_header_style)
            for idx, day in enumerate(days_order)
        ]
        table_data = [header_row]

        for person in people:
            row = [Paragraph(f"<b>{person}</b>", cell_style)]
            for day in days_order:
                match = next((r for r in days[day] if r["person"] == person), None)
                if match:
                    task_clean = strip_emoji(match["task"])
                    hexcolor = task_color.get(task_clean, "#000000")
                    rest_str = "" if match["rest"] == "First" else f'<font size="7" color="#888888"> r{match["rest"]}</font>'
                    row.append(Paragraph(f'<font color="{hexcolor}"><b>{task_clean}</b></font>{rest_str}', cell_style))
                else:
                    row.append(Paragraph("", cell_style))
            table_data.append(row)

        timeline_table = Table(table_data, colWidths=[name_col_width] + [day_col_width] * 7,
                                rowHeights=[26] + [26] * len(people))
        timeline_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#333333")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 1), (0, -1), colors.HexColor("#F2F2F2")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("ROWBACKGROUNDS", (1, 1), (-1, -1), [colors.white, colors.HexColor("#FAFAFA")]),
        ]))
        elements.append(timeline_table)
        elements.append(Spacer(1, 20))

    doc.build(elements)
    return path